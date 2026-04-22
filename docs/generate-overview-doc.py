"""One-off generator for Application-Overview.docx — run: python docs/generate-overview-doc.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Wakalat Namas — Application Overview", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "District Bar Association Bahawalnagar (DBA BWN)"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.italic = True

    doc.add_paragraph()

    doc.add_heading("Why we built this application", level=1)
    doc.add_paragraph(
        "This system supports the controlled issuance of official Wakalat Namas (legal "
        "authorization forms) for the District Bar Association. It exists so the Bar can:"
    )
    bullets_purpose = [
        "Maintain a reliable digital record of how many forms were printed, when, and under which batch.",
        "Assign unique serial numbers to each form so every document can be traced back to an official issuance.",
        "Reduce fraud and casual duplication by tying printing to authenticated staff actions and a central ledger.",
        "Let members of the public verify whether a serial number is genuine—without needing an account.",
    ]
    for b in bullets_purpose:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("What we have implemented (at a glance)", level=1)
    doc.add_paragraph(
        "The product is a web application hosted online, with an optional Windows desktop "
        "wrapper for the same experience where local printing is required."
    )

    features = [
        (
            "Secure access for staff",
            "Staff sign in with email and password. Sessions are managed securely through the backend; sensitive configuration stays on the server.",
        ),
        (
            "Roles and responsibilities",
            "Different roles exist (for example viewers, administrators, and super administrators). "
            "Navigation and sensitive areas (such as user management) are shown only to the right role.",
        ),
        (
            "Dashboard",
            "A central area for authorized users after login—branded for DBABWN and focused on day-to-day operations.",
        ),
        (
            "Printing new batches",
            "Authorized roles can prepare a print run: quantity, calculated fee, and (in the desktop app) choice of physical printer. "
            "Before a batch is reserved in the system, the user must confirm their password again—an extra safeguard if a device is left unattended.",
        ),
        (
            "Official print template",
            "Forms render on a fixed legal-size layout suitable for printers. "
            "The design includes visible and subtle anti-tamper elements: repeating watermark text, "
            "fine decorative patterns, microprinting with serial and validity wording, and a unique "
            "dot pattern—so altered or copied documents are harder to pass off as originals.",
        ),
        (
            "Ledger and verification (staff)",
            "Authorized users can work with batch and form records in the system to support audits and day-to-day checks.",
        ),
        (
            "Public serial verification",
            "The public landing experience allows checking a serial number against the ledger for authenticity, supporting transparency.",
        ),
        (
            "Desktop application (Electron)",
            "For Windows, a desktop build loads the same application from the official deployment and adds integration with installed printers "
            "(and supports workflows intended for office hardware). The app identifies itself so the server can distinguish desktop use from a normal browser when needed.",
        ),
    ]

    for heading, body in features:
        p = doc.add_paragraph()
        p.add_run(f"{heading}. ").bold = True
        p.add_run(body)

    doc.add_heading("How to use this document", level=1)
    doc.add_paragraph(
        "This overview is meant for technical review: it describes the purpose of the system and "
        "the capabilities we have delivered, without going into implementation details. "
        "Your feedback on scope, risks, and possible improvements is welcome."
    )

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Generated for: ").italic = True
    meta.add_run("External / senior engineering review").italic = True

    out = Path(__file__).resolve().parent / "WakalatNamas-Application-Overview.docx"
    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
